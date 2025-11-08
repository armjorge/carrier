
import os
from dotenv import load_dotenv
from colorama import Fore, Style, init
import pandas as pd
from docx import Document
import yaml
from sqlalchemy import create_engine
import subprocess
from datetime import datetime


class CV_GENERATION():
    def postgre_to_docx(self):
        init(autoreset=True)
        print(f"{Fore.BLUE}CARRIER MANAGEMENT{Style.RESET_ALL}")
        self.templates_path = os.path.join(self.working_folder, "CV Templates")
        os.makedirs(self.templates_path, exist_ok=True)

        query = "SELECT * FROM career_accelerator.applications"
        connexion = self.sql_conexion(self.data_access['DB_URL']).connect()
        if connexion is None:
            print("❌ No se pudo establecer conexión con SQL Server.")
            return False

        try:
            self.df_applications = pd.read_sql(query, connexion)
            self.df_cover_letters = pd.read_sql("SELECT * FROM career_accelerator.cover_letters", connexion)
            connexion.close()
            print(f"✅ Loaded applications: {len(self.df_applications)} registros.")
        except Exception as e:
            print(f"❌ Error ejecutando la consulta SQL: {e}")
            return

        df_cv, df_cl = self.get_desired_row(self.df_applications, self.df_cover_letters)
        lang = df_cv['lang'].values[0]
        job = df_cv['job'].values[0]
        cv_path = os.path.join(self.templates_path, f"Curriculum_{lang}.docx")
        cover_letter_path = os.path.join(self.templates_path, f"Cover_letter_{lang}.docx")
        output_cv = os.path.join(self.output_path, f"{job}_CV.docx")
        output_cl = os.path.join(self.output_path, f"{job}_CL.docx")
        if not os.path.exists(cv_path):
            print(f"{Fore.RED}❌ No se encontró el template en: {cv_path}{Style.RESET_ALL}")
            return

        if not os.path.exists(cover_letter_path):
            print(f"{Fore.RED}❌ No se encontró el template en: {cover_letter_path}{Style.RESET_ALL}")
            return
        
        # Selecciona template según modo
        print(f"{Fore.CYAN}📄 Generando currículum...{Style.RESET_ALL}")


        self.populate_document(cv_path, df_cv, output_cv)
        self.open_word_path(output_cv)
        print(f"{Fore.CYAN}📄 Generando carta...{Style.RESET_ALL}")
        
        self.populate_document(cover_letter_path, df_cl, output_cl)
        self.open_word_path(output_cl)
        
    def open_word_path(self, path):
        """Open a file in the default application, cross-platform."""
        if os.name == 'nt':
            os.startfile(path)
        elif os.uname().sysname == 'Darwin':
            subprocess.call(['open', path])
        else:
            subprocess.call(['xdg-open', path])

    def populate_document(self, template_doc, df, output_file):
        for _, row in df.iterrows():
            job = str(row.get("job", "Unknown"))
            try:
                doc = Document(template_doc)

                # 🔹 Reemplazar placeholders en párrafos con soporte para saltos de línea y bullets
                for p in doc.paragraphs:
                    for key in df.columns:
                        placeholder = f"{{{key}}}"
                        if placeholder in p.text:
                            value = str(row.get(key, "")).replace('\\n', '\n')
                            parts = value.split('\n')

                            # Reemplaza el placeholder por la primera línea
                            p.text = p.text.replace(placeholder, parts[0])

                            # Si hay más líneas, las inserta como nuevos párrafos con el mismo estilo
                            if len(parts) > 1:
                                p_element = p._element
                                body_element = doc._body._element
                                index = list(body_element).index(p_element)

                                for part in parts[1:]:
                                    new_p = doc.add_paragraph(part, style=p.style.name)
                                    new_p_element = new_p._element
                                    body_element.remove(new_p_element)
                                    body_element.insert(index + 1, new_p_element)
                                    index += 1
                doc.save(output_file)
                print(f"{Fore.GREEN}✅ Curriculum generado: {output_file}{Style.RESET_ALL}")

            except Exception as e:
                print(f"{Fore.RED}❌ Error generando {job}: {e}{Style.RESET_ALL}")
         
    def get_desired_row(self, df_cv, df_cl):
        columns_pk = ['job', 'lang', 'company_name']
        max_length = len(df_cv)
        for index, row in df_cv.iterrows():
            print(f"{index} - {row[columns_pk]}")
        while True:
            selected_indices = input("Ingrese la fila que requieres para generar el cv")
            try:
                selected_index = int(selected_indices)
                if 0 <= selected_index < max_length:
                    break
                else:
                    print(f"Por favor, ingrese un número entre 0 y {max_length - 1}")
            except ValueError:
                print("Por favor, ingrese un número entero válido")
        selected_row = df_cv.iloc[[selected_index]]
        # 🔹 Filtrar df_cl donde job, lang y company_name coincidan
        match_mask = (
            (df_cl["job"] == selected_row["job"].values[0]) &
            (df_cl["lang"] == selected_row["lang"].values[0]) &
            (df_cl["company_name"] == selected_row["company_name"].values[0])
        )
        df_cl_match = df_cl.loc[match_mask]

        print("Ingresa la fecha que quieras que aparezca en la carta (formato DD/MM/AAAA): \n")
        str_date = input('DD/MM/AAAA: ')
        input_date = datetime.strptime(str_date, '%d/%m/%Y') if str_date else None

        if input_date:
            day = input_date.day
            month_num = input_date.month
            year = input_date.year

            months = {
                'English': ["January", "February", "March", "April", "May", "June",
                            "July", "August", "September", "October", "November", "December"],
                'Spanish': ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"],
                'French': ["janvier", "février", "mars", "avril", "mai", "juin",
                        "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
            }
            lang = df_cv['lang'].values[0]
            if lang == 'English':
                suffix = 'th'
                if day in [1, 21, 31]:
                    suffix = 'st'
                elif day in [2, 22]:
                    suffix = 'nd'
                elif day in [3, 23]:
                    suffix = 'rd'
                date_issued = f"Mexico City, {months['English'][month_num-1]} {day}{suffix}, {year}"

            elif lang == 'Spanish':
                date_issued = f"Ciudad de México, {day} de {months['Spanish'][month_num-1].capitalize()} de {year}"

            elif lang == 'French':
                date_issued = f"Mexico, le {day} {months['French'][month_num-1].capitalize()} {year}"

            else:
                date_issued = input_date.strftime('%d/%m/%Y')
        else:
            date_issued = datetime.today().strftime('%d/%m/%Y')


        # Agregar al DataFrame
        selected_row['date_issued'] = date_issued
        df_cl['date_issued'] = date_issued

        return selected_row, df_cl
    

    def sql_conexion(self, sql_url):
        try:
            engine = create_engine(sql_url)
            return engine
        except Exception as e:
            print(f"❌ Error connecting to database: {e}")
            return None
    # Initialize the main components
    def __init__(self, working_folder, data_access):
        self.working_folder = working_folder
        os.makedirs(self.working_folder, exist_ok=True)
        self.data_access = data_access
        self.output_path = os.path.join(self.working_folder, "Output CVs")
        os.makedirs(self.output_path, exist_ok=True)
        
if __name__ == "__main__":
    env_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    env_file = os.path.join(env_path, '.env')
    folder_name = "MAIN_PATH"
    db_key = "DB_URL"
    
    working_folder = "."
    pg_dict = {}

    if os.path.exists(env_file):
        load_dotenv(dotenv_path=env_file)
        working_folder = os.getenv(folder_name)
        pg_dict = {"DB_URL": os.getenv(db_key)} 

    yaml_path = os.path.join(env_path, 'config', 'config.yml')
    with open(yaml_path, 'r') as file:
        data_access = yaml.safe_load(file)
        if data_access is None:
            data_access = {}
        data_access.update(pg_dict)  # ⚠️ Esto solo funciona si data_access es una lista

    app = CV_GENERATION(working_folder, data_access)
    app.postgre_to_docx()